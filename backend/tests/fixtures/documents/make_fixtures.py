"""Generate binary fixtures (PDF, DOCX, encrypted DOCX) for ingestion tests.

Run manually with: python tests/fixtures/documents/make_fixtures.py
Outputs are committed alongside this script so tests don't require regeneration.
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent


def make_pdf() -> None:
    import fitz  # PyMuPDF

    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Sample Policy Document")
    page1.insert_text((72, 100), "This is page one with some Unicode: café, 中文.")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "Page two paragraph content for extraction tests.")
    doc.save(FIXTURES_DIR / "sample.pdf")
    doc.close()


def _png_chunk(tag: bytes, data: bytes) -> bytes:
    import struct
    import zlib

    return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data))


def _build_minimal_png() -> bytes:
    """Build a valid 1x1 grayscale PNG using zlib for correct IDAT compression."""
    import struct
    import zlib

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 0, 0, 0, 0)  # 1x1, 8-bit grayscale.
    raw_scanline = b"\x00\x00"  # Filter byte 0 + one black pixel.
    idat = zlib.compress(raw_scanline)
    return (
        signature
        + _png_chunk(b"IHDR", ihdr)
        + _png_chunk(b"IDAT", idat)
        + _png_chunk(b"IEND", b"")
    )


def make_image_only_pdf() -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    png_bytes = _build_minimal_png()
    rect = fitz.Rect(36, 36, 200, 200)
    page.insert_image(rect, stream=png_bytes)
    doc.save(FIXTURES_DIR / "image_only.pdf")
    doc.close()


def make_encrypted_pdf() -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Secret content")
    doc.save(
        FIXTURES_DIR / "encrypted.pdf",
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw="owner",
        user_pw="user",
    )
    doc.close()


def make_docx() -> None:
    import docx

    document = docx.Document()
    document.add_heading("Sample Policy", level=1)
    document.add_paragraph("This is the introductory paragraph with Unicode: café, naïve, 中文.")
    document.add_heading("Details Section", level=2)
    document.add_paragraph("A paragraph under the details section.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Value 1"
    table.cell(1, 1).text = "Value 2"
    document.save(FIXTURES_DIR / "sample.docx")


def make_encrypted_docx() -> None:
    """Write a minimal OLE-CFB (compound file) header to mimic an encrypted OOXML file.

    Real encrypted .docx files are OLE-CFB containers, not zip archives, so they
    fail zip parsing. We only need the magic bytes for detection tests.
    """
    ole_signature = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    padding = b"\x00" * 504  # Pad to a plausible minimum sector size (512 bytes).
    (FIXTURES_DIR / "encrypted.docx").write_bytes(ole_signature + padding)


def make_corrupt_docx() -> None:
    (FIXTURES_DIR / "corrupt.docx").write_bytes(b"not a real docx file at all")


if __name__ == "__main__":
    make_pdf()
    make_image_only_pdf()
    make_encrypted_pdf()
    make_docx()
    make_encrypted_docx()
    make_corrupt_docx()
    print("Fixtures written to", FIXTURES_DIR)

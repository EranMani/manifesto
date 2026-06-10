"""Document ingestion — extraction, structural chunking, embedding, and publication.

Pipeline: lock document row -> extract text (PDF/DOCX/TXT/MD) -> normalize ->
structure-then-token chunking -> batched embeddings -> atomic publish (replace
chunks, set chunk_count, status='ready'). On any failure the document is marked
'failed' with a sanitized error and no partial chunks are visible.
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import re
import unicodedata
from dataclasses import dataclass, field
from typing import TYPE_CHECKING

import tiktoken
from sqlalchemy import delete, text

if TYPE_CHECKING:
    from sqlalchemy.ext.asyncio import AsyncSession

    from app.services.llm import EmbeddingService

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class IngestionError(Exception):
    """Base class for all ingestion errors. Messages are sanitized for storage."""


class UnsupportedContentTypeError(IngestionError):
    """Content type is not one of the supported document formats."""


class EmptyDocumentError(IngestionError):
    """Document contains no extractable text."""


class EncryptedDocumentError(IngestionError):
    """Document requires a password to open."""


class CorruptDocumentError(IngestionError):
    """Document bytes could not be parsed as the declared format."""


class ImageOnlyDocumentError(IngestionError):
    """Document contains only images/scans; OCR is out of scope."""


class InvalidEncodingError(IngestionError):
    """TXT/MD content is not valid UTF-8."""


# ---------------------------------------------------------------------------
# Data types
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class ExtractedBlock:
    """One unit of structurally-meaningful text (paragraph, heading, table row)."""

    text: str
    page_number: int | None
    section: str | None


@dataclass(frozen=True)
class ExtractedDocument:
    """Result of format-specific extraction, before normalization/chunking."""

    blocks: list[ExtractedBlock]


@dataclass(frozen=True)
class Chunk:
    """A token-bounded chunk ready for embedding and storage."""

    chunk_index: int
    content: str
    token_count: int
    page_number: int | None
    section: str | None


@dataclass(frozen=True)
class IngestResult:
    """Outcome of an ingestion attempt."""

    document_id: str
    status: str
    chunk_count: int
    error_message: str | None = field(default=None)


# ---------------------------------------------------------------------------
# Chunking configuration
# ---------------------------------------------------------------------------

_TARGET_TOKENS = 450
_HARD_CAP_TOKENS = 600
_OVERLAP_TOKENS = 60

_TOKENIZER_NAME = "cl100k_base"

_MAX_PAGES = 2000
_MAX_BLOCKS = 50_000
_MAX_BLOCK_CHARS = 200_000


def _get_encoding() -> tiktoken.Encoding:
    return tiktoken.get_encoding(_TOKENIZER_NAME)


# ---------------------------------------------------------------------------
# Normalization
# ---------------------------------------------------------------------------

_WHITESPACE_RUN = re.compile(r"[ \t ]+")
_BLANK_LINE_RUN = re.compile(r"\n{3,}")


def _normalize_text(raw: str) -> str:
    """NFC-normalize unicode and collapse redundant whitespace, preserving line breaks."""
    text_ = unicodedata.normalize("NFC", raw)
    lines = [_WHITESPACE_RUN.sub(" ", line).strip() for line in text_.split("\n")]
    text_ = "\n".join(lines)
    text_ = _BLANK_LINE_RUN.sub("\n\n", text_)
    return text_.strip()


# ---------------------------------------------------------------------------
# Format-specific extraction
# ---------------------------------------------------------------------------


def _extract_pdf(file_bytes: bytes) -> ExtractedDocument:
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise CorruptDocumentError("Could not open PDF document.") from exc

    try:
        if doc.is_encrypted:
            raise EncryptedDocumentError("PDF is encrypted and cannot be ingested.")

        if doc.page_count > _MAX_PAGES:
            raise CorruptDocumentError("PDF exceeds the maximum supported page count.")

        blocks: list[ExtractedBlock] = []
        has_text = False
        has_images = False

        for page_index in range(doc.page_count):
            page = doc[page_index]
            page_text = page.get_text("text") or ""
            if page.get_images(full=True):
                has_images = True
            normalized = _normalize_text(page_text)
            if not normalized:
                continue
            has_text = True
            for paragraph in normalized.split("\n\n"):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                if len(blocks) >= _MAX_BLOCKS:
                    raise CorruptDocumentError("PDF exceeds the maximum supported structure size.")
                blocks.append(ExtractedBlock(text=paragraph, page_number=page_index + 1, section=None))

        if not has_text:
            if has_images:
                raise ImageOnlyDocumentError("PDF contains only images; OCR is not supported.")
            raise EmptyDocumentError("PDF contains no extractable text.")

        return ExtractedDocument(blocks=blocks)
    finally:
        doc.close()


_OLE_CFB_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _extract_docx(file_bytes: bytes) -> ExtractedDocument:
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    # Encrypted OOXML files are OLE-CFB containers, not zips; detect before opening as zip.
    if file_bytes.startswith(_OLE_CFB_SIGNATURE):
        raise EncryptedDocumentError("DOCX is encrypted and cannot be ingested.")

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except PackageNotFoundError as exc:
        raise CorruptDocumentError("Could not open DOCX document.") from exc
    except Exception as exc:
        message = str(exc).lower()
        if "encrypt" in message or "password" in message:
            raise EncryptedDocumentError("DOCX is encrypted and cannot be ingested.") from exc
        raise CorruptDocumentError("Could not open DOCX document.") from exc

    blocks: list[ExtractedBlock] = []
    current_section: str | None = None

    for paragraph in document.paragraphs:
        raw = paragraph.text
        normalized = _normalize_text(raw)
        if not normalized:
            continue
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        if style_name.lower().startswith("heading") or style_name.lower() == "title":
            current_section = normalized
        if len(blocks) >= _MAX_BLOCKS:
            raise CorruptDocumentError("DOCX exceeds the maximum supported structure size.")
        blocks.append(ExtractedBlock(text=normalized, page_number=None, section=current_section))

    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if not cells:
                continue
            row_text = " | ".join(cells)
            if len(blocks) >= _MAX_BLOCKS:
                raise CorruptDocumentError("DOCX exceeds the maximum supported structure size.")
            blocks.append(ExtractedBlock(text=row_text, page_number=None, section=current_section))

    if not blocks:
        raise EmptyDocumentError("DOCX contains no extractable text.")

    return ExtractedDocument(blocks=blocks)


def _extract_plain_text(file_bytes: bytes) -> ExtractedDocument:
    try:
        raw = file_bytes.decode("utf-8", errors="strict")
    except UnicodeDecodeError as exc:
        raise InvalidEncodingError("Document is not valid UTF-8 text.") from exc

    if len(raw) > _MAX_BLOCK_CHARS * _MAX_BLOCKS:
        raise CorruptDocumentError("Document exceeds the maximum supported size.")

    normalized = _normalize_text(raw)
    if not normalized:
        raise EmptyDocumentError("Document contains no extractable text.")

    blocks: list[ExtractedBlock] = []
    current_section: str | None = None
    for paragraph in normalized.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        first_line = paragraph.split("\n", 1)[0].strip()
        if first_line.startswith("#"):
            current_section = first_line.lstrip("#").strip() or current_section
        if len(blocks) >= _MAX_BLOCKS:
            raise CorruptDocumentError("Document exceeds the maximum supported structure size.")
        blocks.append(ExtractedBlock(text=paragraph, page_number=None, section=current_section))

    return ExtractedDocument(blocks=blocks)


_EXTRACTORS = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "text/plain": _extract_plain_text,
    "text/markdown": _extract_plain_text,
}


def extract_document(file_bytes: bytes, content_type: str) -> ExtractedDocument:
    """Dispatch to the format-specific extractor. Raises IngestionError subclasses."""
    extractor = _EXTRACTORS.get(content_type)
    if extractor is None:
        raise UnsupportedContentTypeError(f"Unsupported content type: {content_type!r}")
    if not file_bytes:
        raise EmptyDocumentError("Uploaded file is empty.")
    return extractor(file_bytes)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def _split_oversized_block(block: ExtractedBlock, encoding: tiktoken.Encoding) -> list[ExtractedBlock]:
    """Split a single block whose token count exceeds the hard cap into hard-cap-sized pieces."""
    tokens = encoding.encode(block.text)
    if len(tokens) <= _HARD_CAP_TOKENS:
        return [block]
    pieces: list[ExtractedBlock] = []
    for start in range(0, len(tokens), _HARD_CAP_TOKENS):
        piece_tokens = tokens[start : start + _HARD_CAP_TOKENS]
        piece_text = encoding.decode(piece_tokens)
        pieces.append(ExtractedBlock(text=piece_text, page_number=block.page_number, section=block.section))
    return pieces


def chunk_blocks(blocks: list[ExtractedBlock]) -> list[Chunk]:
    """Group structural blocks into token-bounded chunks with overlap within a section.

    Deterministic: same input blocks always produce the same ordered chunks.
    """
    if not blocks:
        return []

    encoding = _get_encoding()

    # Pre-split any block that alone exceeds the hard cap.
    flat_blocks: list[ExtractedBlock] = []
    for block in blocks:
        flat_blocks.extend(_split_oversized_block(block, encoding))

    chunks: list[Chunk] = []
    current_blocks: list[ExtractedBlock] = []
    current_tokens = 0

    def _flush() -> None:
        nonlocal current_blocks, current_tokens
        if not current_blocks:
            return
        content = "\n\n".join(b.text for b in current_blocks)
        page_number = next((b.page_number for b in current_blocks if b.page_number is not None), None)
        section = next((b.section for b in current_blocks if b.section is not None), None)
        token_count = len(encoding.encode(content))
        chunks.append(
            Chunk(
                chunk_index=len(chunks),
                content=content,
                token_count=token_count,
                page_number=page_number,
                section=section,
            )
        )

    previous_section: str | None = None
    for block in flat_blocks:
        block_tokens = len(encoding.encode(block.text))
        section_changed = current_blocks and block.section != previous_section

        would_overflow = current_tokens + block_tokens > _HARD_CAP_TOKENS
        if current_blocks and (would_overflow or section_changed or current_tokens >= _TARGET_TOKENS):
            _flush()
            # Build overlap from the tail of the flushed chunk, only within the same section.
            if not section_changed:
                tail_blocks: list[ExtractedBlock] = []
                tail_token_count = 0
                for prev_block in reversed(current_blocks):
                    prev_tokens = len(encoding.encode(prev_block.text))
                    if tail_token_count + prev_tokens > _OVERLAP_TOKENS:
                        break
                    tail_blocks.insert(0, prev_block)
                    tail_token_count += prev_tokens
                current_blocks = tail_blocks
                current_tokens = tail_token_count
            else:
                current_blocks = []
                current_tokens = 0

        current_blocks.append(block)
        current_tokens += block_tokens
        previous_section = block.section

    _flush()
    return chunks


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------


def _sanitize_error(exc: Exception) -> str:
    """Return a short, safe-to-store error message (no stack traces, no raw bytes)."""
    if isinstance(exc, IngestionError):
        return str(exc)
    return "Document ingestion failed due to an internal error."


def _document_lock_key(document_id: str) -> int:
    """Derive a stable 63-bit advisory lock key from the document id."""
    digest = hashlib.sha256(document_id.encode("utf-8")).digest()
    return int.from_bytes(digest[:8], "big", signed=False) >> 1


async def ingest_document(
    *,
    document_id: str,
    file_bytes: bytes,
    filename: str,
    content_type: str,
    db: AsyncSession,
    embeddings: EmbeddingService,
) -> IngestResult:
    """Extract, chunk, embed, and atomically publish a policy document.

    Acquires a per-document advisory lock for the duration of the ingestion so
    concurrent ingestion attempts on the same document serialize. On failure,
    marks the document 'failed' with a sanitized error and leaves no partial chunks.
    """
    from app.models.policy import PolicyChunk

    lock_key = _document_lock_key(document_id)
    await db.execute(text("SELECT pg_advisory_xact_lock(:key)"), {"key": lock_key})

    try:
        extracted = await asyncio.to_thread(extract_document, file_bytes, content_type)
        chunks = await asyncio.to_thread(chunk_blocks, extracted.blocks)

        if not chunks:
            raise EmptyDocumentError("Document contains no extractable text.")

        contents = [c.content for c in chunks]
        vectors = await embeddings.embed_documents(contents)
        profile = embeddings.profile

        if len(vectors) != len(chunks):
            raise IngestionError("Embedding count does not match chunk count.")

        await db.execute(
            delete(PolicyChunk).where(PolicyChunk.document_id == document_id)
        )

        for chunk, vector in zip(chunks, vectors):
            db.add(
                PolicyChunk(
                    document_id=document_id,
                    chunk_index=chunk.chunk_index,
                    content=chunk.content,
                    embedding=vector,
                    token_count=chunk.token_count,
                    page_number=chunk.page_number,
                    section=chunk.section,
                    metadata_={},
                )
            )

        result = await db.execute(
            text(
                "UPDATE policy_documents SET status='ready', chunk_count=:count, "
                "error_message=NULL, embedding_provider=:provider, "
                "embedding_model=:model, embedding_dimensions=:dimensions, "
                "updated_at=now() WHERE id=:id"
            ),
            {
                "count": len(chunks),
                "provider": profile.provider,
                "model": profile.model,
                "dimensions": profile.dimensions,
                "id": document_id,
            },
        )
        if result.rowcount == 0:
            raise IngestionError("Document row not found during publish.")

        await db.commit()

        return IngestResult(document_id=document_id, status="ready", chunk_count=len(chunks))

    except Exception as exc:
        await db.rollback()
        sanitized = _sanitize_error(exc)
        logger.warning("Ingestion failed for document %s: %s", document_id, sanitized)

        # New transaction: re-acquire the lock, ensure no partial chunks remain, mark failed.
        await db.execute(text("SELECT pg_advisory_xact_lock(:key)"), {"key": lock_key})
        await db.execute(delete(PolicyChunk).where(PolicyChunk.document_id == document_id))
        await db.execute(
            text(
                "UPDATE policy_documents SET status='failed', error_message=:error, "
                "chunk_count=0, updated_at=now() WHERE id=:id"
            ),
            {"error": sanitized, "id": document_id},
        )
        await db.commit()

        return IngestResult(
            document_id=document_id,
            status="failed",
            chunk_count=0,
            error_message=sanitized,
        )
